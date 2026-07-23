"""
生成面向医生/心理治疗师的项目流程说明文档 — Word (.docx)
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ── 页面设置 ──────────────────────────────────────────
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── 样式配置 ──────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5
# Set East Asian font
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Heading styles
for level, (size, color, space_before, space_after, bold) in enumerate([
    (22, RGBColor(0x1B, 0x3A, 0x4B), 24, 12, True),   # Heading 1
    (16, RGBColor(0x0D, 0x7C, 0x85), 18, 8, True),     # Heading 2
    (13, RGBColor(0x1B, 0x3A, 0x4B), 12, 6, True),     # Heading 3
], start=1):
    h_style = doc.styles[f'Heading {level}']
    h_font = h_style.font
    h_font.name = '微软雅黑'
    h_font.size = Pt(size)
    h_font.color.rgb = color
    h_font.bold = bold
    h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h_style.paragraph_format.space_before = Pt(space_before)
    h_style.paragraph_format.space_after = Pt(space_after)
    h_style.paragraph_format.line_spacing = 1.3

# ── 辅助函数 ──────────────────────────────────────────
def add_para(text, style_name='Normal', bold=False, italic=False, size=None, color=None, alignment=None, space_after=None):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    return p

def add_table_with_style(headers, rows, col_widths=None):
    """创建带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 深蓝绿背景
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A4B"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # 数据行
    for i, row_data in enumerate(rows):
        bg = "F5FAF9" if i % 2 == 0 else "FFFFFF"
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = '微软雅黑'
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}"/>')
            cell._tc.get_or_add_tcPr().append(shading)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table

def add_note_box(text, label="说明"):
    """添加说明框"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run_label = p.add_run(f"【{label}】")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = RGBColor(0x0D, 0x7C, 0x85)
    run_text = p.add_run(f" {text}")
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return p

def add_clinical_scenario(title, steps):
    """添加临床场景框"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(f"临床场景：{title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x0D, 0x7C, 0x85)

    for i, step in enumerate(steps, 1):
        bp = doc.add_paragraph()
        bp.paragraph_format.left_indent = Cm(0.8)
        bp.paragraph_format.space_after = Pt(2)
        rn = bp.add_run(f"{i}. {step}")
        rn.font.size = Pt(10.5)
        rn.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # thin line via border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

# ═══════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════

# 空行撑开
for _ in range(4):
    doc.add_paragraph()

# 类型标签
add_para("临床流程说明文档", size=14, color=RGBColor(0x0D, 0x7C, 0x85),
         bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

# 主标题
add_para("心理异常智能早筛与精准干预系统", size=28,
         color=RGBColor(0x1B, 0x3A, 0x4B), bold=True,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

# 副标题
add_para("面向临床心理科与精神科的 AI 辅助工具", size=14,
         color=RGBColor(0x5C, 0xD6, 0xC3), bold=False,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

# 分隔线
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:bottom w:val="single" w:sz="8" w:space="1" w:color="5CD6C3"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)

doc.add_paragraph()

# 日期
today = datetime.date.today().strftime("%Y年%m月%d日")
add_para(today, size=12, color=RGBColor(0x99, 0x99, 0x99),
         alignment=WD_ALIGN_PARAGRAPH.CENTER)

# 分页
doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 目录页
# ═══════════════════════════════════════════════════════
add_heading("目  录", level=1)
toc_items = [
    ("一、系统简介", "什么是这个系统？它能做什么？与医生的关系"),
    ("二、患者交互流程 — 类比门诊", "从患者进入系统到完成干预的完整路径"),
    ("三、核心步骤详解", "安全把关  ·  情绪识别  ·  分诊路由  ·  干预执行"),
    ("四、特殊场景处理", "紧急情况  ·  量表筛查  ·  视频情绪识别"),
    ("五、系统边界说明", "能做什么 / 不能做什么 / 适用场景"),
    ("六、系统在临床中的定位", "总结：这个工具在诊疗流程中扮演什么角色"),
]
for title, desc in toc_items:
    p = doc.add_paragraph()
    run_t = p.add_run(title)
    run_t.bold = True
    run_t.font.size = Pt(13)
    run_t.font.color.rgb = RGBColor(0x1B, 0x3A, 0x4B)
    run_d = p.add_run(f"    {desc}")
    run_d.font.size = Pt(10)
    run_d.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 一、系统简介
# ═══════════════════════════════════════════════════════
add_heading("一、系统简介", level=1)

add_heading("1.1 这是什么系统？", level=2)
add_para(
    "心理异常智能早筛与精准干预系统，是一个面向心理健康领域的 AI 辅助工具。"
    "它像一个\"24小时在线的心理初筛助手\"，能够在医生不在场的情况下，"
    "自动完成对患者的初步评估、情绪分析与干预引导。"
)
add_para(
    "系统的核心设计理念来源于临床诊疗的实际流程：就像医生接诊时会先排查风险、再了解情绪、"
    "再判断严重程度、最后给出处置方案一样，系统也遵循\"安全分诊 → 情绪评估 → 分诊决策 → 干预执行\"的四步流程。"
    "不同的是，它把这个过程自动化了，使得患者在正式见到医生之前，就能获得一次有质量的初步筛查。"
)

add_heading("1.2 它能做什么？", level=2)
items = [
    ("危险信号识别：", "第一时间排查患者表述中的自杀、自伤、暴力等高危内容，一旦发现立即启动危机干预流程。"),
    ("情绪状态评估：", "综合患者的文字、语音、面部表情三个维度，判断当前主导情绪类型和情绪强度。"),
    ("分诊方向判断：", "根据危险等级和情绪严重程度，自动将患者分流到\"危机干预\"\"共情安抚\"或\"知识科普\"三个通道之一。"),
    ("专业内容回复：", "对于非危机患者，系统能生成具备心理学专业水准的对话回复；对于危机患者，系统自动输出规范的安抚话术并推送求助信息。"),
    ("量表自动筛查：", "内置 11 个临床常用心理评估量表，涵盖抑郁（PHQ-9）、焦虑（GAD-7）、躯体症状（PHQ-15）、创伤应激（PCL-5）、睡眠（ISI）、酒精使用（AUDIT）、强迫（OCI-R）、ADHD（ASRS）、社交焦虑（LSAS）、进食障碍（SCOFF）、双相障碍（MDQ），可根据对话内容自动匹配并引导患者完成量表作答，模拟医生问诊收集量表信息的过程。"),
]
for bold_part, text in items:
    add_bullet(text, bold_prefix=bold_part)

add_heading("1.3 和医生工作的关系", level=2)
add_para(
    "本系统定位为辅助工具，而非替代品。其核心价值体现在以下几个场景："
)
add_clinical_scenario("场景一：候诊期间的初筛", [
    "张先生因\"最近总是睡不好、心情低落\"通过线上渠道接触到系统。",
    "在等待正式门诊（可能是几天甚至几周后）的间隙，张先生先和系统进行了约 15 分钟的对话。",
    "系统判断：情绪为\"悲伤\"，强度中等，PHQ-9 自评得分为 14 分（中度抑郁倾向）。",
    "当张先生走进诊室时，医生已经拿到了这份初步筛查记录——\"悲伤情绪，中度，PHQ-9=14\"——作为面诊参考。",
])
add_clinical_scenario("场景二：深夜的危机防护", [
    "凌晨 2 点，一名有抑郁史的患者在群聊或线上平台表达了\"不想活了\"的内容。",
    "此时没有医生在线。系统检测到危险等级为 2 级（高危），立即自动触发危机响应：",
    "——向患者推送规范的心理安抚话术；",
    "——提示求助渠道（24小时心理援助热线：400-161-9995）；",
    "——记录本次危机事件并标记为紧急。",
    "次日，相关临床人员可查阅系统记录的危机事件，安排后续随访。",
])

add_note_box("系统的输出信息提供的是辅助参考，不是诊断结论。最终诊断、治疗方案的确定，必须由具备执业资质的临床医生完成。", "重要提醒")

# ═══════════════════════════════════════════════════════
# 二、患者交互流程 — 类比门诊
# ═══════════════════════════════════════════════════════
doc.add_page_break()
add_heading("二、患者交互流程 — 类比门诊流程", level=1)

add_para(
    "为了让临床医生更容易理解系统的运作方式，下面将系统的四步流程类比成大家熟悉的门诊工作流程："
)

add_para("", space_after=4)

# 流程类比表
flow = [
    ("患者走进诊室", "→", "患者打开系统界面",
     "患者可以通过打字、语音说话、甚至录制一段短视频来与系统交流。\n系统界面类似聊天窗口，操作门槛极低。"),
    ("第一步：安全分诊\n（排查危险）", "→", "安全过滤\n（Safety Filter）",
     "就像在急诊科，分诊台护士最先要判断\"这个患者有没有生命危险\"。\n系统也会在第一时间扫描患者的每一句话，看是否包含自杀、自残、暴力等高危内容。"),
    ("第二步：情绪问诊\n（了解情绪状态）", "→", "情绪分析\n（Emotion Analysis）",
     "就像医生通过望、闻、问来判断患者的情绪状态。\n系统通过三个\"感官\"来感知：看文字（患者写了什么）、听语音（声音中的情绪）、看面部（视频中的表情）。"),
    ("第三步：分诊决策\n（判断处理方向）", "→", "分诊路由\n（Route Decision）",
     "就像门诊分诊后决定\"这个患者该去哪个科室\"。\n系统综合危险等级和情绪分析结果，自动判断应该走\"危机处置\"\"心理安抚\"还是\"健康教育\"通道。"),
    ("第四步：干预执行\n（给出处置方案）", "→", "干预生成\n（Intervention）",
     "就像医生开出处方或给出建议。\n系统根据分诊结果，走不同的干预通道，给患者一个有针对性的回复。"),
]

table = doc.add_table(rows=len(flow) + 1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
# 列宽
for row in table.rows:
    row.cells[0].width = Cm(3.2)
    row.cells[1].width = Cm(0.8)
    row.cells[2].width = Cm(3.2)
    row.cells[3].width = Cm(6.8)

# 表头
header_labels = ["门诊类比", "", "系统流程", "详细说明"]
for j, label in enumerate(header_labels):
    cell = table.rows[0].cells[j]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = '微软雅黑'
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A4B"/>')
    cell._tc.get_or_add_tcPr().append(shading)

for i, (clinical, arrow, tech, desc) in enumerate(flow):
    bg = "F5FAF9" if i % 2 == 0 else "FFFFFF"
    for j_val, (val, bold, sz, clr, align) in enumerate([
        (clinical, True, 10, RGBColor(0x1B, 0x3A, 0x4B), WD_ALIGN_PARAGRAPH.CENTER),
        (arrow, False, 16, RGBColor(0x0D, 0x7C, 0x85), WD_ALIGN_PARAGRAPH.CENTER),
        (tech, True, 10, RGBColor(0x0D, 0x7C, 0x85), WD_ALIGN_PARAGRAPH.CENTER),
        (desc, False, 9.5, RGBColor(0x55, 0x55, 0x55), WD_ALIGN_PARAGRAPH.LEFT),
    ]):
        cell = table.rows[i + 1].cells[j_val]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.bold = bold
        run.font.size = Pt(sz)
        run.font.name = '微软雅黑'
        run.font.color.rgb = clr
        p.alignment = align
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

doc.add_paragraph()

add_note_box("上表左侧是医生日常工作中的环节，右侧是系统对应的自动化步骤。整个流程的设计逻辑，就是为了模拟一个规范的临床初筛过程。", "阅读提示")

# ═══════════════════════════════════════════════════════
# 三、核心步骤详解
# ═══════════════════════════════════════════════════════
doc.add_page_break()
add_heading("三、核心步骤详解", level=1)

# --- 步骤一 ---
add_heading("步骤一：安全把关 — 第一时间排查危机", level=2)

add_para("▎解决了什么临床问题？", bold=True, size=11, color=RGBColor(0x1B, 0x3A, 0x4B))
add_para(
    "在临床实践中，最需要优先处理的情况就是患者表达了自杀、自伤或伤人的倾向。"
    "但在线上场景中，医生不可能 7×24 小时守在屏幕前。安全把关这一步就是用系统替代人工，"
    "在患者发出危险信号的第一时间就做出响应——\"先保安全，再谈诊疗\"。"
)

add_para("▎它是怎么工作的？", bold=True, size=11, color=RGBColor(0x1B, 0x3A, 0x4B))

steps_data = [
    ("患者输入", "患者在聊天框输入文字，或者通过语音说话，或者上传视频。系统接收到的所有内容都会首先进入这个安全过滤环节。"),
    ("关键词匹配", "系统内置了一个包含 10000 余条中文敏感词的词库，这些词汇按危险程度分为 0、1、2 三个等级。\n• Level 0：一般敏感词（如\"烦躁\"\"失眠\"）\n• Level 1：较重敏感词（如\"活不下去\"\"想消失\"）\n• Level 2：极端危险词（如\"我要自杀\"\"再见了世界\"等明确的自我伤害意图表达）"),
    ("等级判定", "当检测到 Level 2 词汇时，系统不做任何等待，立即将整个流程转入紧急状态。\n这就像急诊分诊台的\"红色标签\"——一旦挂上，走的是完全不同的处理通道。"),
    ("输出结果", "系统输出一个明确的安全评估结论：\n• 是否有危险\n• 危险等级（0/1/2）\n• 具体匹配到了哪些敏感词\n• 是否需要阻断或进入紧急流程"),
]

for title, desc in steps_data:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"   {title}  ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor(0x0D, 0x7C, 0x85)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

add_separator()

# --- 步骤二 ---
add_heading("步骤二：情绪识别 — 多维度了解患者情绪状态", level=2)

add_para("▎解决了什么临床问题？", bold=True, size=11, color=RGBColor(0x1B, 0x3A, 0x4B))
add_para(
    "医生在面诊时，判断患者情绪状态依赖多个信息来源——患者说了什么（问诊内容）、怎么说（语调语气）、"
    "表情如何（面部观察）。系统通过 AI 技术，模拟了这个\"望闻问\"的多维度观察过程。"
    "对于线上问诊场景中医生无法面对面观察患者的痛点，系统提供了额外的信息补充。"
)

add_para("▎三个维度的情绪感知", bold=True, size=11, color=RGBColor(0x1B, 0x3A, 0x4B))

add_table_with_style(
    ["感知维度", "对应医生工作", "系统具体做法", "输出信息"],
    [
        ["文字情绪", "问诊 — 听患者说了什么", "分析患者输入的文字内容，识别七种基本情绪类型（愤怒、厌恶、恐惧、快乐、悲伤、惊讶、中性）并评估情绪强度", "主导情绪类型 + 情绪强度评分"],
        ["语音情绪", "听诊 — 听患者的语气", "从语音中提取声学特征（语调高低、语速快慢、停顿模式），判断潜藏在声音里的情绪线索", "语音情绪分类 + 置信度"],
        ["面部情绪", "望诊 — 观察表情", "分析视频中患者的面部表情变化，逐帧识别微表情，捕捉文字和声音可能隐瞒的真实情绪", "面部情绪分类 + 逐帧情绪变化序列"],
    ],
    col_widths=[2.2, 3.0, 5.0, 3.8]
)

add_para("▎多模态情绪融合", bold=True, size=11, color=RGBColor(0x1B, 0x3A, 0x4B))
add_para(
    "三个维度的情绪信息会被送入\"融合引擎\"进行综合判断。融合的原则是："
)
add_bullet("当三个维度的判断一致时（比如文字、语音、面部都指向\"悲伤\"），输出结论的置信度最高。")
add_bullet("当信号之间存在冲突时（比如文字说\"我没事\"但面部表情显示悲伤），系统会按各通道的置信度加权裁决——通常面部表情和语音信号的权重会高于纯文字，因为非语言信号更难伪装。")
add_bullet("如果缺少某个维度（比如患者只发了文字，没有语音或视频），系统会基于已有的信息给出判断，但会在输出中标注\"缺少某些维度的信息\"。")

add_para(
    "最终，系统输出一个统一的情绪评估结论，包含：住院情绪类型、情绪强度（0-10）、风险等级、"
    "以及备注说明（如\"多信号冲突，以面部信号为准\"），供后续的分诊环节使用。"
)

add_separator()

# --- 步骤三 ---
add_heading("步骤三：分诊路由 — 决定患者走哪条干预通道", level=2)

add_para("▎解决了什么临床问题？", bold=True, size=11, color=RGBColor(0x1B, 0x3A, 0x4B))
add_para(
    "在精神科或心理科的日常工作中，医生在初步了解患者情况后，会迅速做一个\"分流判断\"："
    "这个患者是需要紧急处理（比如有自杀倾向），还是需要重点安抚（情绪波动大），"
    "还是可以通过心理教育和科普来帮助（情况较轻或只是来咨询的）？"
    "分诊路由这一步，就是用规则系统来模拟这个判断过程。"
)

add_para("▎三条分诊通道", bold=True, size=11, color=RGBColor(0x1B, 0x3A, 0x4B))

add_table_with_style(
    ["通道", "适用情况", "处理原则", "类比临床场景"],
    [
        ["危机通道\n（Crisis）",
         "安全等级达到 Level 2\n或情绪强度极高（≥ 0.8）\n且伴随高风险信号",
         "最高优先级，不经过 AI 生成\n直接使用经过审核的固定安抚模板\n确保回复的确定性、稳定性",
         "急诊红色通道\n立即响应，不上门诊排队"],
        ["安抚通道\n（Comfort）",
         "情绪强度中等（0.4~0.8）\n或主导情绪为负面情绪\n（悲伤、恐惧、愤怒等）",
         "以共情和安抚为核心\n回应要温暖、专业、有支持感\n可采用多种心理治疗技术",
         "普通门诊\n以情绪疏导为主"],
        ["知识通道\n（Knowledge）",
         "情绪强度较低（< 0.4）\n或主导情绪偏中性/正面\n或患者主动寻求知识",
         "以心理健康教育和信息科普为主\n可结合知识库给出建议\n可启动量表自评流程",
         "健康教育咨询\n提供信息和工具"],
    ],
    col_widths=[2.2, 3.8, 4.5, 3.5]
)

add_para("▎决策逻辑说明", bold=True, size=11, color=RGBColor(0x1B, 0x3A, 0x4B))
add_para(
    "路由决策并非\"非此即彼\"，而是一个带置信度的判断。系统会给每个决策一个 0~1 的置信度分数。"
    "置信度越高，说明信息越充分、判断越可靠；置信度偏低时，说明不同信号之间存在矛盾，"
    "或者患者表达的信息不够明确。路由模块还会计算出\"这条判断距离另一条通道的边界有多近\"——"
    "如果靠近边界，说明患者可能同时需要两种通道的干预。"
)

add_separator()

# --- 步骤四 ---
add_heading("步骤四：干预执行 — 根据分诊结果给出针对性回复", level=2)

add_para("▎解决了什么临床问题？", bold=True, size=11, color=RGBColor(0x1B, 0x3A, 0x4B))
add_para(
    "分诊之后，不同通道的患者需要不同形式的回应。就像急诊科、普通门诊和健康咨询中心"
    "的工作方式完全不同，系统的三个通道也采用了截然不同的回复生成策略。"
)

add_para("▎三条通道的干预方式", bold=True, size=11, color=RGBColor(0x1B, 0x3A, 0x4B))

add_heading("通道 A：危机干预 — 确定性优先", level=3)
add_bullet("回复策略：使用预先编写并经过临床审核的固定安抚模板，不依赖 AI 生成（因为在危机场景下，AI 生成的内容可能存在不可预测的输出风险）。")
add_bullet("安抚话术内容：包括共情陈述、安全确认、求助引导。例如：\"我听到了你的痛苦，你现在的感受是真实且重要的。请知道，你不必一个人面对这一切...\"")
add_bullet("同时触发紧急推送：系统会在后台记录本次危机事件，并推送求助资源（心理援助热线、附近精神科急诊信息等）。")
add_bullet("确定性保障：模板内容不会因为输入变化而产生不可预期的输出——每次同样的危机信号触发同样经过审核的安抚话术。")

add_heading("通道 B：安抚干预 — 共情与支持", level=3)
add_bullet("回复策略：调用大语言模型生成共情性对话，融入专业心理治疗技术。")
add_bullet("可调用的八种心理治疗技术：倾听与共情、认知重构、行为激活、放松训练、接纳承诺疗法（ACT）、积极心理学技术、自我关怀、社交支持激活。")
add_bullet("系统会根据当前的对话上下文，选择最合适的技术方向来引导对话。例如，当患者表达\"我觉得自己一无是处\"时，系统可能选择\"认知重构\"方向，帮助患者审视自动化负性思维。")

add_heading("通道 C：知识干预 — 科普教育与量表筛查", level=3)
add_bullet("回复策略：结合心理健康知识库进行精准信息检索，生成有据可依的科普性回复。")
add_bullet("知识库涵盖的内容类型：抑郁、焦虑、应激障碍、睡眠问题、人际关系、情绪管理等常见心理议题。")
add_bullet("如果患者尚未做过量表筛查，系统可根据对话主题自动匹配 11 个量表中合适的量表（如谈到情绪低落可触发 PHQ-9，谈到睡眠问题可触发 ISI），引导患者完成自评流程，就像医生在问诊中顺带让患者做个筛查问卷一样。")

add_note_box(
    "三条通道并非完全隔离。在实际应用中，一个患者的多次对话可能在不同通道之间切换——"
    "比如，某次对话走安抚通道进行共情疏导，下一次如果情绪稳定、开始询问知识问题，"
    "就可能切换到知识通道。系统会根据每次对话的内容实时做出新的路由判断。",
    "通道切换说明"
)

# ═══════════════════════════════════════════════════════
# 四、特殊场景处理
# ═══════════════════════════════════════════════════════
doc.add_page_break()
add_heading("四、特殊场景处理", level=1)

add_heading("4.1 紧急情况处理流程", level=2)
add_para(
    "当系统在安全过滤环节检测到危险等级 ≥ 2（即患者表达了明确的自我伤害或伤害他人的意图），"
    "整个处理流程会发生以下变化："
)

add_para("", space_after=2)
add_bullet("跳过环节：系统将跳过\"情绪分析\"和\"分诊路由\"两个步骤，直接进入危机干预通道。")
add_bullet("自动赋值：情绪自动标记为\"极度痛苦（distress）\"，路由自动设为\"危机通道\"。")
add_bullet("固定回复：不经过 AI 模型生成内容，直接使用预先审核的危机安抚模板。")
add_bullet("紧急推送：在回复中明确呈现求助热线和紧急联系方式。")
add_bullet("后台记录：系统将本次事件记录为\"紧急标记\"，供后续临床随访使用。")
add_para("")

add_para(
    "设计考量：在危机场景下，\"确定性\"比\"智能性\"更重要。必须保证系统给患者的回复是经过"
    "临床审核的、安全可靠的、不会出现任何意外或不当表述的内容。任何 AI 生成的不可控性在危机场景下都不可接受。"
)
add_note_box("相当于急诊分诊制度中的\"先抢救、后挂号\"原则。安全永远是第一优先级。", "临床类比")

add_heading("4.2 量表自动筛查（11 个临床量表）", level=2)
add_para(
    "系统内置了 11 个临床常用心理评估量表，覆盖抑郁、焦虑、躯体症状、创伤应激、睡眠、"
    "物质使用、强迫、注意缺陷、社交焦虑、进食障碍和双相障碍等常见心理健康维度的筛查需求。"
    "所有量表均按标准题序和计分规则实现，能够像医生逐题询问患者一样，引导患者完成完整的量表作答。"
)

add_para("内置量表清单：", bold=True, size=11, color=RGBColor(0x1B, 0x3A, 0x4B))
add_table_with_style(
    ["量表", "中文名称", "题目数", "筛查维度"],
    [
        ["PHQ-9", "患者健康问卷抑郁量表", "9", "近两周抑郁症状频率"],
        ["GAD-7", "广泛性焦虑障碍量表", "7", "近两周焦虑症状频率"],
        ["PHQ-15", "患者健康问卷躯体症状量表", "15", "躯体不适症状困扰程度"],
        ["PCL-5", "创伤后应激障碍筛查量表", "20", "创伤事件后应激症状"],
        ["ISI", "失眠严重程度指数量表", "7", "入睡、维持睡眠、早醒等问题"],
        ["AUDIT", "酒精使用障碍筛查量表", "10", "酒精使用量与依赖风险"],
        ["OCI-R", "强迫症状量表修订版", "18", "强迫思维与强迫行为"],
        ["ASRS", "成人 ADHD 自评量表", "6", "注意力缺陷与多动冲动"],
        ["LSAS", "Liebowitz 社交焦虑量表", "24", "社交场景中的恐惧与回避"],
        ["SCOFF", "进食障碍筛查量表", "5", "厌食症与贪食症快速筛查"],
        ["MDQ", "双相障碍筛查问卷", "15", "躁狂/轻躁狂发作史"],
    ],
    col_widths=[1.5, 4.2, 1.3, 7.0]
)

add_para("量表筛查流程：", bold=True, size=11, color=RGBColor(0x1B, 0x3A, 0x4B))
add_table_with_style(
    ["步骤", "系统行为", "类比医生的操作"],
    [
        ["智能匹配", "根据对话主题自动从 11 个量表中选出最合适的量表（如患者谈到睡眠问题→ISI，谈到饮酒→AUDIT）", "医生判断\"现在需要做哪个筛查？\""],
        ["启动判断", "系统评估当前对话状态：患者是否适合现在做量表？是否已经在做某个量表中途？", "医生判断\"现在是否需要让患者做量表？\""],
        ["逐题施测", "按照量表标准题序逐一呈现，每道题等待患者回答后再出下一题", "医生口头逐题询问"],
        ["作答追踪", "记录每道题的回答，维护答题进度；中断后下次可以继续", "医生在病历本上记录各题分数"],
        ["评分输出", "全部题目完成后，计算总分并给出分级提示（含临床阈值说明）", "医生计算总分、评估严重程度"],
        ["结果反馈", "将量表评分以通俗语言反馈给患者，同时作为后续 AI 回复的参考背景", "医生向患者解释评估结果并讨论后续方案"],
    ],
    col_widths=[2.0, 5.5, 6.5]
)

add_note_box(
    "所有量表筛查结果仅供辅助参考，不作为诊断依据。各量表均为自评筛查工具，"
    "阳性结果提示需要进一步临床评估，不能等同于诊断结论。医生应当结合面诊综合判断。",
    "临床注意事项"
)

add_heading("4.3 视频情绪识别", level=2)
add_para(
    "除了文字和语音，系统还支持患者通过录制短视频来进行情绪表达和分析。"
    "这一功能在以下场景中特别有用："
)
add_bullet("当患者的文字表达能力有限（如不愿打字、文化程度限制等），但愿意通过视频交流时。")
add_bullet("当患者说\"我没事\"但临床经验提示可能有情绪掩饰时——面部微表情和声音特征能提供额外线索。")
add_bullet("远程问诊场景中，视频情绪识别相当于为线上诊疗增加了一个\"望诊\"维度。")
add_para("")

add_para("视频分析的工作方式：", bold=True, size=11, color=RGBColor(0x1B, 0x3A, 0x4B))
add_bullet("提取音频 → 转换为文字（语音识别）+ 分析声音情绪。")
add_bullet("提取面部帧 → 逐帧分析面部表情类型和强度变化。")
add_bullet("综合三路信号 → 输出统一的情绪评估结论。")
add_bullet("需要至少 3 个有效面部帧 + 人脸检出率 ≥30% 才会启动面部情绪分析（避免因遮挡、光线等原因导致的不可靠判断）。")

# ═══════════════════════════════════════════════════════
# 五、系统边界说明
# ═══════════════════════════════════════════════════════
doc.add_page_break()
add_heading("五、系统边界说明", level=1)

add_heading("5.1 系统能做什么", level=2)
items_can = [
    "对患者的文字、语音、视频输入进行初步的风险筛查，识别危险信号。",
    "从多个维度评估患者的情绪状态，给出情绪类型和强度的参考判断。",
    "根据危险等级和情绪状态，自动将患者分流到合适的干预通道。",
    "对非危机患者，生成具有一定专业水准的共情性对话或心理健康科普回复。",
    "引导患者完成 11 个标准化心理评估量表的自评（PHQ-9、GAD-7、PHQ-15、PCL-5、ISI、AUDIT、OCI-R、ASRS、LSAS、SCOFF、MDQ），自动计分和分级。",
    "在检测到高危情况时，自动触发危机干预模板并推送求助信息。",
    "24 小时运行，不受时间、人力限制。",
    "记录和分析每次对话的内容与情绪变化轨迹，为后续的临床判断提供参考数据。",
]
for item in items_can:
    add_bullet(item)

add_heading("5.2 系统不能做什么", level=2)
items_cannot = [
    ("不能替代临床诊断：", "系统的所有输出都是\"分析结果\"和\"辅助建议\"，不是诊断结论。精神科诊断需要医生综合问诊、体检、辅助检查等信息，绝不是仅靠文字分析就能完成的。"),
    ("没有处方权：", "系统不能开具处方、不能推荐药物、不能进行任何需要执业医师资格的操作。"),
    ("不能处理复杂临床情况：", "对于共病、器质性疾病引起的精神症状、需要鉴别诊断的复杂病例，系统不具备处理能力，必须由临床医生介入。"),
    ("不能替代治疗关系：", "治疗联盟（Therapeutic Alliance）的建立需要真正的人际互动。AI 可以提供支持性对话，但不能替代真实的治疗关系。"),
    ("不保证 100% 准确：", "情绪分析、路由决策等环节基于概率模型和规则引擎，存在一定误差。最终的判断需要医生确认。"),
]
for bold_part, text in items_cannot:
    add_bullet(text, bold_prefix=bold_part)

add_heading("5.3 建议的使用场景", level=2)
add_para("本系统最适合在以下场景中发挥作用：")
add_bullet("基层医疗机构：作为心理问题初筛工具，帮助全科医生或社区医生在面对心理健康主诉时有辅助参考。")
add_bullet("学校心理咨询中心：作为学生的第一接触点，完成初步评估，缓解咨询师人力不足的压力。")
add_bullet("企业员工心理援助（EAP）：提供 24 小时可用的自助心理支持，在员工求助后的第一时间给出响应。")
add_bullet("候诊/等待期的过渡支持：在患者等待正式心理治疗或精神科门诊期间，提供低强度的过渡性支持。")
add_bullet("远程医疗/互联网医院的心理卫生模块：作为线上第一层筛查，辅助远程医生更高效地了解患者状态。")
add_bullet("大规模心理普查活动：在面对大量受检人群时，提供标准化的初筛流程，标记出需要重点关注的对象。")

# ═══════════════════════════════════════════════════════
# 六、系统在临床中的定位
# ═══════════════════════════════════════════════════════
doc.add_page_break()
add_heading("六、系统在临床中的定位", level=1)

add_para(
    "如果用一个比喻来说明这个系统在临床诊疗流程中的位置——"
)
add_para("")

add_para(
    "它就像一个\"智能分诊台\"。", bold=True, size=14, color=RGBColor(0x0D, 0x7C, 0x85),
    alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para("")

add_para(
    "在患者走进医生诊室之前，分诊台完成了一轮有质量的初步筛查。它让紧急的患者被快速识别和响应，"
    "让情绪困扰的患者得到第一时间的共情和支持，让情况较轻或求信息的患者获得结构化的科普与量表工具。"
    "而所有这些筛查过程中的记录——患者说了什么、情绪如何波动、量表得分是多少——都整理好，交到医生手中，"
    "让后续的临床诊疗更高效、更有针对性。"
)

add_para("")

add_para("系统的核心设计原则——", bold=True, size=12, color=RGBColor(0x1B, 0x3A, 0x4B))

principles = [
    ("安全优先：", "所有对话以安全过滤为第一步。发现危险即触发紧急响应，不走常规流程。"),
    ("辅助而非替代：", "永远作为辅助工具存在。不做诊断、不开处方、不替代医生判断。"),
    ("多维度评估：", "综合文字、语音、面部三个维度的信息，尽可能全面地了解患者状态。"),
    ("场景适配性：", "根据不同场景（紧急/安抚/教育）提供不同形式和策略的干预。"),
    ("持续可追踪：", "每次对话和情绪变化都有记录，形成可追踪的患者画像，为临床决策提供参考。"),
]
for bold_part, text in principles:
    add_bullet(text, bold_prefix=bold_part)

doc.add_paragraph()
add_separator()
doc.add_paragraph()

add_para("", space_after=8)
add_para(
    "本系统目前仍在持续研发和优化中。我们期待与临床一线的医生和心理治疗师紧密合作，"
    "根据实际使用中的反馈不断完善功能，让这个工具真正成为心理健康领域的有效辅助手段。",
    size=11, color=RGBColor(0x66, 0x66, 0x66), alignment=WD_ALIGN_PARAGRAPH.CENTER
)

# ── 页脚 ──────────────────────────────────────────────
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("心理异常智能早筛与精准干预系统 · 临床流程说明文档")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.font.name = '微软雅黑'

# ── 保存 ──────────────────────────────────────────────
output_path = r"c:\Users\xt\Desktop\Project\mental-intervene-master\项目流程说明文档（临床版）.docx"
doc.save(output_path)
print(f"文档已生成: {output_path}")
